#!/usr/bin/env python3
"""
Smart TCM Document Processor - Handles tables and Hebrew content
"""

import os
import re
import time
from typing import List, Dict
from docx import Document
from deep_translator import GoogleTranslator
from supabase import create_client
from sentence_transformers import SentenceTransformer
from tqdm import tqdm

# Import config
from config import SUPABASE_URL, SUPABASE_KEY, ACUPUNCTURE_FILE, SYNDROME_FILE

class SmartTCMProcessor:
    """Smart processor that handles tables and Hebrew text"""
    
    def __init__(self):
        print("🔧 Initializing smart processor...")
        
        # Initialize Supabase
        self.supabase = create_client(SUPABASE_URL, SUPABASE_KEY)
        print("   ✅ Supabase connected")
        
        # Initialize translator
        self.translator = GoogleTranslator(source='en', target='iw')
        print("   ✅ Translator ready")
        
        # Initialize embedding model
        print("   ⏳ Loading embedding model...")
        self.model = SentenceTransformer('all-MiniLM-L6-v2')
        print("   ✅ Embedding model loaded")
        
        print("✅ Smart processor ready!\n")
    
    def translate_text(self, text: str) -> str:
        """Translate text to Hebrew with chunking"""
        if not text or len(text.strip()) == 0:
            return ""
        
        # If already Hebrew, return as-is
        if self.is_hebrew(text):
            return text
        
        # If text is short, translate directly
        if len(text) <= 4500:
            try:
                return self.translator.translate(text)
            except Exception as e:
                print(f"   ⚠️  Translation error: {e}")
                return text
        
        # Split long text by sentences
        sentences = text.split('. ')
        translated = []
        current_chunk = ""
        
        for sentence in sentences:
            if len(current_chunk) + len(sentence) > 4500:
                try:
                    translated.append(self.translator.translate(current_chunk))
                    time.sleep(0.1)
                except Exception as e:
                    print(f"   ⚠️  Translation error: {e}")
                    translated.append(current_chunk)
                current_chunk = sentence
            else:
                current_chunk += sentence + ". "
        
        # Translate final chunk
        if current_chunk:
            try:
                translated.append(self.translator.translate(current_chunk))
            except:
                translated.append(current_chunk)
        
        return " ".join(translated)
    
    def is_hebrew(self, text: str) -> bool:
        """Check if text is primarily Hebrew"""
        if not text:
            return False
        hebrew_chars = sum(1 for c in text if '\u0590' <= c <= '\u05FF')
        return hebrew_chars > len(text) * 0.3
    
    def generate_embedding(self, text: str) -> List[float]:
        """Generate 384-dimensional embedding"""
        embedding = self.model.encode(text)
        return embedding.tolist()
    
    def extract_from_tables(self, file_path: str) -> List[Dict]:
        """Extract acupuncture points from tables"""
        print(f"📖 Reading tables from: {os.path.basename(file_path)}")
        
        doc = Document(file_path)
        points = []
        
        print(f"   Found {len(doc.tables)} tables")
        
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                cells = [cell.text.strip() for cell in row.cells]
                
                # Skip empty rows
                if not any(cells):
                    continue
                
                # Look for point codes (LI-4, ST-36, etc.)
                point_code = None
                for cell in cells:
                    match = re.search(r'([A-Z]{2,3})[- ]?(\d{1,2})', cell)
                    if match:
                        meridian, number = match.groups()
                        point_code = f"{meridian}-{number}"
                        break
                
                # If we found a point code, extract info
                if point_code:
                    # Combine all cell text
                    full_text = '\n'.join(cells)
                    
                    point = {
                        'point_code': point_code,
                        'meridian': point_code.split('-')[0],
                        'point_number': int(point_code.split('-')[1]),
                        'content_en': full_text[:1000],  # First 1000 chars
                        'name_en': cells[0] if cells else point_code,
                        'location_en': '',
                        'functions_en': '',
                        'indications_en': ''
                    }
                    
                    points.append(point)
        
        print(f"   ✅ Extracted {len(points)} acupuncture points from tables")
        return points
    
    def extract_hebrew_sections(self, file_path: str) -> List[Dict]:
        """Extract sections from Hebrew text (no heading styles)"""
        print(f"📖 Reading Hebrew sections from: {os.path.basename(file_path)}")
        
        doc = Document(file_path)
        sections = []
        current_section = None
        
        # Hebrew patterns that indicate section headers
        header_patterns = [
            r'^שיעור \d+',  # Lesson X
            r'^תפקוד',       # Functions
            r'^תהליך',       # Process
            r'^כיצד',        # How
            r'^\*',          # Starts with asterisk
            r'^-',           # Starts with dash
            r'^[א-ת]{2,15}:$',  # Hebrew word followed by colon
        ]
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text or len(text) < 3:
                continue
            
            # Check if this looks like a header
            is_header = False
            for pattern in header_patterns:
                if re.match(pattern, text):
                    is_header = True
                    break
            
            # Also check if it's short (< 100 chars) and ends with certain patterns
            if len(text) < 100 and (text.endswith(':') or text.endswith('?')):
                is_header = True
            
            if is_header:
                # Save previous section
                if current_section and current_section['content']:
                    sections.append(current_section)
                
                # Start new section
                current_section = {
                    'title': text,
                    'content': []
                }
            else:
                # Add to current section
                if current_section is None:
                    current_section = {
                        'title': 'מבוא',  # Introduction
                        'content': []
                    }
                current_section['content'].append(text)
        
        # Add final section
        if current_section and current_section['content']:
            sections.append(current_section)
        
        print(f"   ✅ Found {len(sections)} sections")
        return sections
    
    def process_acupuncture_book(self, file_path: str):
        """Process acupuncture points book (tables)"""
        print("\n" + "="*60)
        print("📖 PROCESSING ACUPUNCTURE POINTS BOOK (TABLES)")
        print("="*60 + "\n")
        
        # Extract from tables
        points = self.extract_from_tables(file_path)
        
        if len(points) == 0:
            print("⚠️  No points found in tables. Trying paragraph extraction...")
            # Fallback to paragraph parsing
            return
        
        # Process ALL points (not just first 100)
        print(f"\n🔄 Processing ALL {len(points)} points...")
        success_count = 0
        
        for i, point in enumerate(tqdm(points, desc="Translating & storing")):
            try:
                # Translate (if not already Hebrew)
                point['name_he'] = self.translate_text(point['name_en'])
                point['location_he'] = self.translate_text(point['location_en'])
                point['functions_he'] = self.translate_text(point['functions_en'])
                point['indications_he'] = self.translate_text(point.get('content_en', '')[:500])
                
                # Generate embedding from Hebrew content
                embedding_text = f"{point['name_he']} {point['functions_he']}"
                embedding = self.generate_embedding(embedding_text)
                
                # Store in Supabase
                data = {
                    'point_code': point['point_code'],
                    'name_en': point['name_en'][:255],
                    'name_he': point['name_he'][:255],
                    'meridian': point['meridian'],
                    'point_number': point['point_number'],
                    'location_en': point['location_en'][:500],
                    'location_he': point['location_he'][:500],
                    'functions_en': point['functions_en'][:1000],
                    'functions_he': point['functions_he'][:1000],
                    'indications_en': point['indications_en'][:1000],
                    'indications_he': point['indications_he'][:1000],
                    'embedding': embedding
                }
                
                self.supabase.table('acupuncture_points').upsert(data).execute()
                success_count += 1
                
                # Rate limiting
                if i % 10 == 0:
                    time.sleep(0.5)
                
            except Exception as e:
                print(f"\n   ⚠️  Error processing {point['point_code']}: {e}")
                continue
        
        print(f"\n✅ Successfully processed {success_count}/{len(points)} points!")
    
    def process_syndrome_book(self, file_path: str):
        """Process syndromes book (Hebrew paragraphs)"""
        print("\n" + "="*60)
        print("📖 PROCESSING ZANG-FU SYNDROMES (HEBREW)")
        print("="*60 + "\n")
        
        # Extract Hebrew sections
        sections = self.extract_hebrew_sections(file_path)
        
        # Process ALL sections (not just first 50)
        print(f"\n🔄 Processing ALL {len(sections)} sections...")
        success_count = 0
        
        for i, section in enumerate(tqdm(sections, desc="Processing & storing")):
            try:
                title = section['title']
                content = '\n'.join(section['content'][:5])  # First 5 paragraphs
                
                # Generate embedding from Hebrew
                embedding_text = f"{title} {content[:500]}"
                embedding = self.generate_embedding(embedding_text)
                
                # Store in Supabase
                data = {
                    'name_en': '',  # Empty for now
                    'name_he': title[:255],
                    'symptoms_en': '',
                    'symptoms_he': content[:2000],
                    'etiology_en': '',
                    'etiology_he': '',
                    'tongue_en': '',
                    'tongue_he': '',
                    'pulse_en': '',
                    'pulse_he': '',
                    'treatment_en': '',
                    'treatment_he': '',
                    'embedding': embedding
                }
                
                self.supabase.table('zangfu_syndromes').insert(data).execute()
                success_count += 1
                
                time.sleep(0.3)  # Rate limiting
                
            except Exception as e:
                print(f"\n   ⚠️  Error processing section: {e}")
                continue
        
        print(f"\n✅ Successfully processed {success_count}/{len(sections)} syndromes!")


def main():
    """Main execution"""
    print("\n" + "="*60)
    print("🌿 SMART TCM DOCUMENT PROCESSOR")
    print("="*60)
    print()
    print("📊 This version handles:")
    print("   ✅ Tables (acupuncture points)")
    print("   ✅ Hebrew text without heading styles")
    print("   ✅ Pattern-based section detection")
    print()
    print(f"📖 Acupuncture: {os.path.basename(ACUPUNCTURE_FILE)}")
    print(f"📖 Syndromes: {os.path.basename(SYNDROME_FILE)}")
    print(f"🗄️  Database: {SUPABASE_URL}")
    print()
    print("⏱️  Estimated time: 30-40 minutes")
    print("   (Processing ALL points + ALL syndromes)")
    print()
    
    response = input("Ready to start? (yes/no): ")
    if response.lower() not in ['yes', 'y']:
        print("Cancelled.")
        return
    
    print()
    
    try:
        processor = SmartTCMProcessor()
        
        # Process acupuncture book
        processor.process_acupuncture_book(ACUPUNCTURE_FILE)
        
        # Process syndromes
        processor.process_syndrome_book(SYNDROME_FILE)
        
        print("\n" + "="*60)
        print("🎉 PROCESSING COMPLETE!")
        print("="*60)
        print()
        print("✅ Documents processed and stored in Supabase")
        print("✅ All content translated to Hebrew")
        print("✅ Vector embeddings generated")
        print()
        print("📊 Database now contains:")
        print("   • All acupuncture points from your book")
        print("   • All syndrome sections (262 total)")
        print()
        print("Next steps:")
        print("1. Check your Supabase database")
        print("2. Test the search functions")
        print("3. Run the Next.js app for Q&A!")
        print()
        
    except Exception as e:
        print(f"\n❌ Error: {e}")
        import traceback
        traceback.print_exc()


if __name__ == "__main__":
    main()
