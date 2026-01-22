#!/usr/bin/env python3
"""
Document Structure Analyzer
Run this to understand why only 1 section was found
"""

from docx import Document
from config import ACUPUNCTURE_FILE, SYNDROME_FILE

def analyze_document(file_path, doc_name):
    """Analyze document structure"""
    print("\n" + "="*70)
    print(f"{doc_name.upper()}")
    print("="*70)
    
    doc = Document(file_path)
    
    print(f"\n📊 Total paragraphs: {len(doc.paragraphs)}")
    
    # Count by style
    styles = {}
    for para in doc.paragraphs:
        style = para.style.name
        styles[style] = styles.get(style, 0) + 1
    
    print(f"\n📋 Paragraph styles:")
    for style, count in sorted(styles.items(), key=lambda x: -x[1]):
        print(f"   {style}: {count}")
    
    # Show first 30 paragraphs
    print(f"\n📖 First 30 paragraphs (with text):")
    count = 0
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            count += 1
            text = para.text[:80].replace('\n', ' ')
            print(f"   {i+1:3d}. [{para.style.name:20s}] {text}...")
            if count >= 30:
                break
    
    # Check tables
    print(f"\n📊 Tables: {len(doc.tables)}")
    if len(doc.tables) > 0:
        print(f"   First table has {len(doc.tables[0].rows)} rows")

def main():
    print("\n🔍 TCM DOCUMENT STRUCTURE ANALYZER")
    print("This will help us understand why only 1 section was found\n")
    
    try:
        analyze_document(ACUPUNCTURE_FILE, "Acupuncture Points Book")
    except Exception as e:
        print(f"❌ Error analyzing acupuncture book: {e}")
    
    try:
        analyze_document(SYNDROME_FILE, "Syndromes Book")
    except Exception as e:
        print(f"❌ Error analyzing syndromes book: {e}")
    
    print("\n" + "="*70)
    print("💡 RECOMMENDATIONS:")
    print("="*70)
    print("1. Look for the main heading styles used")
    print("2. Check if content is in tables vs paragraphs")
    print("3. Identify point codes (like LI-4, ST-36, etc.)")
    print("\nShare this output so we can fix the parser!")

if __name__ == "__main__":
    main()
