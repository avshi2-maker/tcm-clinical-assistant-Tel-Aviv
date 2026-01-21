#!/usr/bin/env python3
"""
TCM Document Processor - Hebrew Translation & Supabase Integration
Processes TCM Word documents, translates to Hebrew, and creates searchable database
"""

import os
import json
import re
from pathlib import Path
from typing import List, Dict, Optional
import asyncio
from dataclasses import dataclass, asdict

# Document processing
from docx import Document
import pypandoc

# Translation
from deep_translator import GoogleTranslator

# Database
from supabase import create_client, Client

# Embeddings
import openai
from sentence_transformers import SentenceTransformer

# Progress tracking
from tqdm import tqdm


@dataclass
class TCMDocument:
    """Represents a processed TCM document section"""
    id: str
    source_file: str
    section_type: str  # 'point', 'syndrome', 'chapter', etc.
    title_en: str
    title_he: str
    content_en: str
    content_he: str
    metadata: Dict
    embedding: Optional[List[float]] = None


class HebrewTranslator:
    """Handles translation to Hebrew with chunking and rate limiting"""
    
    def __init__(self, batch_size: int = 500):
        self.translator = GoogleTranslator(source='en', target='he')
        self.batch_size = batch_size
        
    def translate_text(self, text: str, max_length: int = 4500) -> str:
        """
        Translate text to Hebrew with automatic chunking
        Google Translate has 5000 char limit per request
        """
        if not text or not text.strip():
            return ""
        
        # If text is short, translate directly
        if len(text) <= max_length:
            try:
                return self.translator.translate(text)
            except Exception as e:
                print(f"Translation error: {e}")
                return text
        
        # Split by sentences for better translation quality
        sentences = re.split(r'([.!?]\s+)', text)
        translated_chunks = []
        current_chunk = ""
        
        for i, sentence in enumerate(sentences):
            # Check if adding this sentence would exceed limit
            if len(current_chunk) + len(sentence) > max_length:
                # Translate current chunk
                if current_chunk:
                    try:
                        translated = self.translator.translate(current_chunk)
                        translated_chunks.append(translated)
                        # Small delay to avoid rate limiting
                        asyncio.sleep(0.1)
                    except Exception as e:
                        print(f"Chunk translation error: {e}")
                        translated_chunks.append(current_chunk)
                    current_chunk = sentence
            else:
                current_chunk += sentence
        
        # Translate final chunk
        if current_chunk:
            try:
                translated = self.translator.translate(current_chunk)
                translated_chunks.append(translated)
            except Exception as e:
                print(f"Final chunk translation error: {e}")
                translated_chunks.append(current_chunk)
        
        return " ".join(translated_chunks)
    
    def translate_batch(self, texts: List[str]) -> List[str]:
        """Translate multiple texts with progress tracking"""
        translated = []
        for text in tqdm(texts, desc="Translating to Hebrew"):
            translated.append(self.translate_text(text))
        return translated


class DocumentProcessor:
    """Extract and structure content from DOCX files"""
    
    @staticmethod
    def extract_from_docx(file_path: str) -> Dict:
        """
        Extract structured content from Word document
        Returns: Dict with sections, paragraphs, and metadata
        """
        try:
            doc = Document(file_path)
            
            sections = []
            current_section = {
                'title': '',
                'content': [],
                'level': 0
            }
            
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                
                # Detect headings by style
                if para.style.name.startswith('Heading'):
                    # Save previous section
                    if current_section['content']:
                        sections.append(current_section.copy())
                    
                    # Start new section
                    level = int(para.style.name.replace('Heading ', ''))
                    current_section = {
                        'title': text,
                        'content': [],
                        'level': level
                    }
                else:
                    current_section['content'].append(text)
            
            # Add final section
            if current_section['content']:
                sections.append(current_section)
            
            return {
                'filename': Path(file_path).name,
                'total_paragraphs': len(doc.paragraphs),
                'sections': sections,
                'metadata': {
                    'core_properties': {
                        'title': doc.core_properties.title or '',
                        'author': doc.core_properties.author or '',
                        'subject': doc.core_properties.subject or ''
                    }
                }
            }
            
        except Exception as e:
            print(f"Error processing {file_path}: {e}")
            return {'filename': Path(file_path).name, 'sections': [], 'error': str(e)}
    
    @staticmethod
    def extract_acupuncture_points(sections: List[Dict]) -> List[Dict]:
        """
        Parse acupuncture point information from sections
        Expected format: Point name, location, functions, indications
        """
        points = []
        
        for section in sections:
            # Try to identify point sections
            title = section['title']
            content = '\n'.join(section['content'])
            
            # Common patterns for acupuncture points
            # e.g., "LI-4 (Hegu)", "ST-36", "SP-6", etc.
            point_pattern = r'([A-Z]{2,3})-?(\d+)'
            matches = re.findall(point_pattern, title)
            
            if matches:
                meridian, number = matches[0]
                point_code = f"{meridian}-{number}"
                
                # Parse content sections
                point_data = {
                    'code': point_code,
                    'name': title,
                    'meridian': meridian,
                    'number': number,
                    'content': content,
                    'location': '',
                    'functions': '',
                    'indications': '',
                    'raw_section': section
                }
                
                # Try to extract specific sections
                content_lower = content.lower()
                
                # Location
                loc_match = re.search(r'location:?\s*([^\n]+)', content_lower)
                if loc_match:
                    point_data['location'] = loc_match.group(1).strip()
                
                # Functions
                func_match = re.search(r'function[s]?:?\s*([^\n]+(?:\n(?!location|indication)[^\n]+)*)', 
                                      content_lower, re.IGNORECASE)
                if func_match:
                    point_data['functions'] = func_match.group(1).strip()
                
                # Indications
                ind_match = re.search(r'indication[s]?:?\s*([^\n]+(?:\n(?!location|function)[^\n]+)*)', 
                                     content_lower, re.IGNORECASE)
                if ind_match:
                    point_data['indications'] = ind_match.group(1).strip()
                
                points.append(point_data)
        
        return points
    
    @staticmethod
    def extract_syndromes(sections: List[Dict]) -> List[Dict]:
        """
        Parse Zang-Fu syndrome information
        Expected format: Syndrome name, etiology, symptoms, treatment
        """
        syndromes = []
        
        for section in sections:
            title = section['title']
            content = '\n'.join(section['content'])
            
            # Syndrome data structure
            syndrome_data = {
                'name': title,
                'content': content,
                'etiology': '',
                'symptoms': '',
                'tongue': '',
                'pulse': '',
                'treatment': '',
                'raw_section': section
            }
            
            # Try to extract specific sections (Hebrew-aware)
            # Look for common section markers in both English and Hebrew
            
            # Etiology / אטיולוגיה
            etio_match = re.search(r'(?:etiology|אטיולוגיה):?\s*([^\n]+(?:\n(?!symptom|treatment|tongue|pulse)[^\n]+)*)', 
                                  content, re.IGNORECASE)
            if etio_match:
                syndrome_data['etiology'] = etio_match.group(1).strip()
            
            # Symptoms / תסמינים
            symp_match = re.search(r'(?:symptom[s]?|תסמינים):?\s*([^\n]+(?:\n(?!etiology|treatment|tongue|pulse)[^\n]+)*)', 
                                  content, re.IGNORECASE)
            if symp_match:
                syndrome_data['symptoms'] = symp_match.group(1).strip()
            
            # Tongue / לשון
            tongue_match = re.search(r'(?:tongue|לשון):?\s*([^\n]+)', content, re.IGNORECASE)
            if tongue_match:
                syndrome_data['tongue'] = tongue_match.group(1).strip()
            
            # Pulse / דופק
            pulse_match = re.search(r'(?:pulse|דופק):?\s*([^\n]+)', content, re.IGNORECASE)
            if pulse_match:
                syndrome_data['pulse'] = pulse_match.group(1).strip()
            
            # Treatment / טיפול
            treat_match = re.search(r'(?:treatment|טיפול):?\s*([^\n]+(?:\n(?!etiology|symptom|tongue|pulse)[^\n]+)*)', 
                                   content, re.IGNORECASE)
            if treat_match:
                syndrome_data['treatment'] = treat_match.group(1).strip()
            
            syndromes.append(syndrome_data)
        
        return syndromes


class SupabaseManager:
    """Manages Supabase database operations for TCM documents"""
    
    def __init__(self, url: str, key: str):
        self.client: Client = create_client(url, key)
        self.embedding_model = SentenceTransformer('all-MiniLM-L6-v2')
    
    async def setup_tables(self):
        """Create necessary database tables and indexes"""
        
        # SQL for creating tables
        create_tables_sql = """
        -- Main TCM documents table
        CREATE TABLE IF NOT EXISTS tcm_documents (
            id UUID PRIMARY KEY DEFAULT gen_random_uuid(),
            source_file TEXT NOT NULL,
            section_type TEXT NOT NULL,
            title_en TEXT,
            title_he TEXT,
            content_en TEXT,
            content_he TEXT,
            metadata JSONB,
            created_at TIMESTAMP WITH TIME ZONE DEFAULT NOW(),
            updated_at TIMESTAMP WITH TIME ZONE DEFAULT NOW()
        );
        
        -- Acupuncture points table
        CREATE TABLE IF NOT EXISTS acupuncture_points (
            id UUID PRIMARY KEY DEFAULT gen_random_uuid(),
            point_code TEXT UNIQUE NOT NULL,
            name_en TEXT,
            name_he TEXT,
            meridian TEXT,
            point_number INTEGER,
            location_en TEXT,
            location_he TEXT,
            functions_en TEXT,
            functions_he TEXT,
            indications_en TEXT,
            indications_he TEXT,
            metadata JSONB,
            embedding vector(384),
            created_at TIMESTAMP WITH TIME ZONE DEFAULT NOW()
        );
        
        -- Zang-Fu syndromes table
        CREATE TABLE IF NOT EXISTS zangfu_syndromes (
            id UUID PRIMARY KEY DEFAULT gen_random_uuid(),
            name_en TEXT,
            name_he TEXT,
            etiology_en TEXT,
            etiology_he TEXT,
            symptoms_en TEXT,
            symptoms_he TEXT,
            tongue_en TEXT,
            tongue_he TEXT,
            pulse_en TEXT,
            pulse_he TEXT,
            treatment_en TEXT,
            treatment_he TEXT,
            metadata JSONB,
            embedding vector(384),
            created_at TIMESTAMP WITH TIME ZONE DEFAULT NOW()
        );
        
        -- Create indexes for faster searches
        CREATE INDEX IF NOT EXISTS idx_tcm_documents_type ON tcm_documents(section_type);
        CREATE INDEX IF NOT EXISTS idx_acupuncture_points_meridian ON acupuncture_points(meridian);
        CREATE INDEX IF NOT EXISTS idx_acupuncture_points_code ON acupuncture_points(point_code);
        
        -- Create vector similarity search functions
        CREATE OR REPLACE FUNCTION match_acupuncture_points(
            query_embedding vector(384),
            match_threshold float,
            match_count int
        )
        RETURNS TABLE (
            id uuid,
            point_code text,
            name_he text,
            similarity float
        )
        LANGUAGE sql STABLE
        AS $$
            SELECT
                id,
                point_code,
                name_he,
                1 - (embedding <=> query_embedding) as similarity
            FROM acupuncture_points
            WHERE 1 - (embedding <=> query_embedding) > match_threshold
            ORDER BY similarity DESC
            LIMIT match_count;
        $$;
        
        CREATE OR REPLACE FUNCTION match_syndromes(
            query_embedding vector(384),
            match_threshold float,
            match_count int
        )
        RETURNS TABLE (
            id uuid,
            name_he text,
            symptoms_he text,
            similarity float
        )
        LANGUAGE sql STABLE
        AS $$
            SELECT
                id,
                name_he,
                symptoms_he,
                1 - (embedding <=> query_embedding) as similarity
            FROM zangfu_syndromes
            WHERE 1 - (embedding <=> query_embedding) > match_threshold
            ORDER BY similarity DESC
            LIMIT match_count;
        $$;
        """
        
        print("Creating database tables...")
        # Note: You'll need to run this SQL in Supabase SQL editor
        # Or use migrations
        return create_tables_sql
    
    def generate_embedding(self, text: str) -> List[float]:
        """Generate embedding vector for text using sentence-transformers"""
        embedding = self.embedding_model.encode(text)
        return embedding.tolist()
    
    async def insert_acupuncture_point(self, point_data: Dict) -> bool:
        """Insert acupuncture point into database"""
        try:
            # Generate embedding from Hebrew content
            embedding_text = f"{point_data['name_he']} {point_data['location_he']} {point_data['functions_he']}"
            embedding = self.generate_embedding(embedding_text)
            
            data = {
                'point_code': point_data['code'],
                'name_en': point_data['name'],
                'name_he': point_data.get('name_he', ''),
                'meridian': point_data['meridian'],
                'point_number': int(point_data['number']),
                'location_en': point_data['location'],
                'location_he': point_data.get('location_he', ''),
                'functions_en': point_data['functions'],
                'functions_he': point_data.get('functions_he', ''),
                'indications_en': point_data['indications'],
                'indications_he': point_data.get('indications_he', ''),
                'metadata': point_data.get('metadata', {}),
                'embedding': embedding
            }
            
            result = self.client.table('acupuncture_points').upsert(data).execute()
            return True
            
        except Exception as e:
            print(f"Error inserting point {point_data.get('code')}: {e}")
            return False
    
    async def insert_syndrome(self, syndrome_data: Dict) -> bool:
        """Insert Zang-Fu syndrome into database"""
        try:
            # Generate embedding from Hebrew content
            embedding_text = f"{syndrome_data['name_he']} {syndrome_data['symptoms_he']} {syndrome_data['treatment_he']}"
            embedding = self.generate_embedding(embedding_text)
            
            data = {
                'name_en': syndrome_data['name'],
                'name_he': syndrome_data.get('name_he', ''),
                'etiology_en': syndrome_data['etiology'],
                'etiology_he': syndrome_data.get('etiology_he', ''),
                'symptoms_en': syndrome_data['symptoms'],
                'symptoms_he': syndrome_data.get('symptoms_he', ''),
                'tongue_en': syndrome_data['tongue'],
                'tongue_he': syndrome_data.get('tongue_he', ''),
                'pulse_en': syndrome_data['pulse'],
                'pulse_he': syndrome_data.get('pulse_he', ''),
                'treatment_en': syndrome_data['treatment'],
                'treatment_he': syndrome_data.get('treatment_he', ''),
                'metadata': syndrome_data.get('metadata', {}),
                'embedding': embedding
            }
            
            result = self.client.table('zangfu_syndromes').insert(data).execute()
            return True
            
        except Exception as e:
            print(f"Error inserting syndrome {syndrome_data.get('name')}: {e}")
            return False
    
    async def semantic_search(self, query: str, table: str = 'acupuncture_points', 
                             limit: int = 5, threshold: float = 0.7) -> List[Dict]:
        """
        Perform semantic search using embeddings
        Args:
            query: Hebrew search query
            table: 'acupuncture_points' or 'zangfu_syndromes'
            limit: Number of results
            threshold: Minimum similarity score
        """
        try:
            # Generate embedding for query
            query_embedding = self.generate_embedding(query)
            
            # Call appropriate RPC function
            if table == 'acupuncture_points':
                result = self.client.rpc('match_acupuncture_points', {
                    'query_embedding': query_embedding,
                    'match_threshold': threshold,
                    'match_count': limit
                }).execute()
            else:
                result = self.client.rpc('match_syndromes', {
                    'query_embedding': query_embedding,
                    'match_threshold': threshold,
                    'match_count': limit
                }).execute()
            
            return result.data
            
        except Exception as e:
            print(f"Search error: {e}")
            return []


class TCMProcessor:
    """Main processor orchestrating the entire pipeline"""
    
    def __init__(self, supabase_url: str, supabase_key: str):
        self.translator = HebrewTranslator()
        self.processor = DocumentProcessor()
        self.db = SupabaseManager(supabase_url, supabase_key)
    
    async def process_acupuncture_book(self, file_path: str):
        """Process acupuncture points book"""
        print(f"\n📖 Processing Acupuncture Points Book: {file_path}")
        
        # Extract content
        doc_data = self.processor.extract_from_docx(file_path)
        print(f"Found {len(doc_data['sections'])} sections")
        
        # Extract points
        points = self.processor.extract_acupuncture_points(doc_data['sections'])
        print(f"Extracted {len(points)} acupuncture points")
        
        # Translate and store
        for i, point in enumerate(tqdm(points, desc="Processing points"), 1):
            # Translate all fields
            point['name_he'] = self.translator.translate_text(point['name'])
            point['location_he'] = self.translator.translate_text(point['location'])
            point['functions_he'] = self.translator.translate_text(point['functions'])
            point['indications_he'] = self.translator.translate_text(point['indications'])
            
            # Store in database
            await self.db.insert_acupuncture_point(point)
            
            # Save progress periodically
            if i % 50 == 0:
                print(f"Processed {i}/{len(points)} points")
        
        print(f"✅ Completed processing {len(points)} acupuncture points")
    
    async def process_syndrome_book(self, file_path: str):
        """Process Zang-Fu syndromes book"""
        print(f"\n📖 Processing Zang-Fu Syndromes: {file_path}")
        
        # Extract content
        doc_data = self.processor.extract_from_docx(file_path)
        print(f"Found {len(doc_data['sections'])} sections")
        
        # Extract syndromes
        syndromes = self.processor.extract_syndromes(doc_data['sections'])
        print(f"Extracted {len(syndromes)} syndromes")
        
        # Translate and store
        for i, syndrome in enumerate(tqdm(syndromes, desc="Processing syndromes"), 1):
            # Translate all fields
            syndrome['name_he'] = self.translator.translate_text(syndrome['name'])
            syndrome['etiology_he'] = self.translator.translate_text(syndrome['etiology'])
            syndrome['symptoms_he'] = self.translator.translate_text(syndrome['symptoms'])
            syndrome['tongue_he'] = self.translator.translate_text(syndrome['tongue'])
            syndrome['pulse_he'] = self.translator.translate_text(syndrome['pulse'])
            syndrome['treatment_he'] = self.translator.translate_text(syndrome['treatment'])
            
            # Store in database
            await self.db.insert_syndrome(syndrome)
            
            if i % 20 == 0:
                print(f"Processed {i}/{len(syndromes)} syndromes")
        
        print(f"✅ Completed processing {len(syndromes)} syndromes")
    
    async def process_all_documents(self, acupuncture_file: str, syndrome_file: str):
        """Process all TCM documents"""
        print("=" * 60)
        print("🌿 TCM Document Processing Pipeline")
        print("=" * 60)
        
        # Process acupuncture points
        await self.process_acupuncture_book(acupuncture_file)
        
        # Process syndromes
        await self.process_syndrome_book(syndrome_file)
        
        print("\n" + "=" * 60)
        print("✅ All documents processed successfully!")
        print("=" * 60)


# ============================================================================
# USAGE EXAMPLES
# ============================================================================

async def main():
    """Main execution function"""
    
    # Configuration
    SUPABASE_URL = os.getenv('SUPABASE_URL', 'your-project-url.supabase.co')
    SUPABASE_KEY = os.getenv('SUPABASE_KEY', 'your-anon-key')
    
    # File paths
    ACUPUNCTURE_FILE = 'ACUPUNCTURE POINTS BOOK^MPOINTS CATACORIES.docx'
    SYNDROME_FILE = 'zang-fu-syndromes-hebrew.docx'
    
    # Initialize processor
    processor = TCMProcessor(SUPABASE_URL, SUPABASE_KEY)
    
    # Setup database (first time only)
    # sql_commands = await processor.db.setup_tables()
    # print("Run these SQL commands in Supabase SQL editor:")
    # print(sql_commands)
    
    # Process all documents
    await processor.process_all_documents(ACUPUNCTURE_FILE, SYNDROME_FILE)
    
    # Example: Search for points
    print("\n🔍 Example Search - Finding points for headache:")
    results = await processor.db.semantic_search(
        query="כאב ראש",  # "headache" in Hebrew
        table='acupuncture_points',
        limit=5
    )
    
    for result in results:
        print(f"  • {result['point_code']}: {result['name_he']} (similarity: {result['similarity']:.2f})")
    
    # Example: Search for syndromes
    print("\n🔍 Example Search - Finding syndromes with fatigue:")
    results = await processor.db.semantic_search(
        query="עייפות",  # "fatigue" in Hebrew
        table='zangfu_syndromes',
        limit=5
    )
    
    for result in results:
        print(f"  • {result['name_he']} (similarity: {result['similarity']:.2f})")


if __name__ == "__main__":
    asyncio.run(main())
